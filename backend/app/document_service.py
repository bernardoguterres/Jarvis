"""Phase 9: explicit browser upload/import of local documents only — never
a scanned or watched folder. Validates content by sniffing actual bytes
(never trusting the filename extension), enforces size/count/page limits,
guards against DOCX zip bombs and macro-enabled documents, extracts text
via safe, non-executing parsers (pypdf/python-docx — no macro or embedded
OLE-object execution), chunks it, and indexes it in the rebuildable FTS5
document index.
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.backup_service import BackupError, create_backup
from app.config import Settings
from app.document_fts_service import remove_document_fts, upsert_document_chunk_fts
from app.recall_index_service import remove_recall, sync_recall
from app.models import Domain
from app.models_integrations import Document, DocumentChunk

MAX_FILES_PER_UPLOAD = 5
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # 25 MiB
MAX_TOTAL_DOCUMENTS = 500
MAX_PDF_PAGES = 500
MAX_DOCX_MEMBERS = 2000
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024  # 100 MiB — zip-bomb guard

CHUNK_SIZE_CHARS = 1500
CHUNK_OVERLAP_CHARS = 200

_SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._-]")


class DocumentValidationError(Exception):
    pass


class DocumentNotFoundError(Exception):
    pass


class DocumentPermanentDeletionError(Exception):
    pass


@dataclass
class SniffResult:
    kind: str  # "pdf" | "docx" | "text"
    mime_type: str


def _sniff(data: bytes, original_filename: str) -> SniffResult:
    if data.startswith(b"%PDF-"):
        return SniffResult("pdf", "application/pdf")

    if data[:2] == b"PK":
        try:
            with zipfile.ZipFile(BytesIO(data)) as zf:
                names = zf.namelist()
        except zipfile.BadZipFile as exc:
            raise DocumentValidationError(f"File claims to be a ZIP/DOCX container but is corrupt: {exc}") from exc
        if "word/document.xml" in names:
            if "word/vbaProject.bin" in names:
                raise DocumentValidationError(
                    "This document contains macros (word/vbaProject.bin) and is not supported."
                )
            return SniffResult(
                "docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        raise DocumentValidationError("File is a ZIP archive but not a recognized .docx document.")

    # No strong binary signature — only accept if it decodes cleanly as text
    # with no embedded NUL bytes (a cheap, effective binary-content check).
    if b"\x00" in data[:8192]:
        raise DocumentValidationError("File content is not recognized as PDF, DOCX, or plain text.")
    try:
        data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentValidationError(f"File is not valid UTF-8 text: {exc}") from exc

    lower_name = original_filename.lower()
    mime = "text/markdown" if lower_name.endswith((".md", ".markdown")) else "text/plain"
    return SniffResult("text", mime)


def _safe_stored_filename(sha256: str, original_filename: str) -> str:
    suffix = Path(original_filename).suffix.lower()
    suffix = _SAFE_FILENAME_RE.sub("", suffix)[:10] or ".bin"
    return f"{sha256}{suffix}"


def _extract_pdf(data: bytes) -> list[tuple[int, str]]:
    from pypdf import PdfReader
    from pypdf.errors import FileNotDecryptedError, PdfReadError

    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise DocumentValidationError(f"Could not parse PDF: {exc}") from exc

    if reader.is_encrypted:
        try:
            result = reader.decrypt("")
            if result == 0:
                raise DocumentValidationError("PDF is encrypted/password-protected — cannot extract text.")
        except (FileNotDecryptedError, PdfReadError, NotImplementedError) as exc:
            raise DocumentValidationError("PDF is encrypted/password-protected — cannot extract text.") from exc

    if len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentValidationError(f"PDF has {len(reader.pages)} pages, exceeding the {MAX_PDF_PAGES}-page limit.")

    pages: list[tuple[int, str]] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append((i + 1, text))
    return pages


def _extract_docx(data: bytes) -> str:
    import docx

    with zipfile.ZipFile(BytesIO(data)) as zf:
        infolist = zf.infolist()
        if len(infolist) > MAX_DOCX_MEMBERS:
            raise DocumentValidationError(f"DOCX contains too many internal parts ({len(infolist)}).")
        total_uncompressed = sum(info.file_size for info in infolist)
        if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise DocumentValidationError("DOCX uncompressed content exceeds the safety limit (possible zip bomb).")

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:
        raise DocumentValidationError(f"Could not parse DOCX: {exc}") from exc

    return "\n".join(p.text for p in document.paragraphs if p.text)


def _build_chunks(pages: list[tuple[int | None, str]]) -> list[tuple[int, int | None, str]]:
    """Fixed-size chunking with overlap, page number carried through when
    the source is paginated (PDF); None for DOCX/TXT/MD."""
    result: list[tuple[int, int | None, str]] = []
    index = 0
    for page_number, text in pages:
        text = text.strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            end = start + CHUNK_SIZE_CHARS
            result.append((index, page_number, text[start:end]))
            index += 1
            if end >= len(text):
                break
            start = end - CHUNK_OVERLAP_CHARS
    return result


def import_document(
    session: Session, settings: Settings, *, domain_id: str, original_filename: str, data: bytes
) -> Document:
    if session.get(Domain, domain_id) is None:
        raise DocumentValidationError(f"Unknown domain_id: {domain_id!r}")
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise DocumentValidationError(f"File exceeds the {MAX_FILE_SIZE_BYTES // (1024*1024)} MiB size limit.")

    total_docs = session.execute(select(Document)).scalars().all()
    if len(total_docs) >= MAX_TOTAL_DOCUMENTS:
        raise DocumentValidationError(f"Document limit ({MAX_TOTAL_DOCUMENTS}) reached.")

    sha256 = hashlib.sha256(data).hexdigest()
    existing = session.execute(select(Document).where(Document.sha256 == sha256)).scalar_one_or_none()
    if existing is not None:
        raise DocumentValidationError(f"Duplicate content — already imported as {existing.original_filename!r}.")

    sniff = _sniff(data, original_filename)

    stored_filename = _safe_stored_filename(sha256, original_filename)
    documents_dir = settings.documents_dir
    documents_dir.mkdir(parents=True, exist_ok=True)
    stored_path = (documents_dir / stored_filename).resolve()
    if documents_dir.resolve() not in stored_path.parents:
        raise DocumentValidationError("Refusing to store outside the documents directory.")

    document = Document(
        domain_id=domain_id,
        original_filename=original_filename[:300],
        stored_relative_path=stored_filename,
        mime_type=sniff.mime_type,
        sha256=sha256,
        size_bytes=len(data),
        status="processing",
    )
    session.add(document)
    session.flush()

    try:
        if sniff.kind == "pdf":
            pages = _extract_pdf(data)
            document.page_count = len(pages)
        elif sniff.kind == "docx":
            pages = [(None, _extract_docx(data))]
        else:
            pages = [(None, data.decode("utf-8"))]

        stored_path.write_bytes(data)

        chunk_tuples = _build_chunks(pages)
        for index, page_number, content in chunk_tuples:
            chunk = DocumentChunk(document_id=document.id, chunk_index=index, page_number=page_number, content=content)
            session.add(chunk)
            session.flush()
            upsert_document_chunk_fts(session, chunk, domain_id)

        document.status = "ready"
    except DocumentValidationError as exc:
        document.status = "encrypted" if "encrypted" in str(exc).lower() else "unsupported"
        document.error_detail = str(exc)[:500]
    except Exception as exc:
        document.status = "error"
        document.error_detail = str(exc)[:500]

    sync_recall(session, "document", document.id)
    session.commit()
    session.refresh(document)
    return document


def get_document_or_404(session: Session, document_id: str) -> Document:
    document = session.get(Document, document_id)
    if document is None:
        raise DocumentNotFoundError(document_id)
    return document


def list_documents(session: Session, *, domain_id: str | None = None) -> list[Document]:
    stmt = select(Document)
    if domain_id is not None:
        stmt = stmt.where(Document.domain_id == domain_id)
    stmt = stmt.order_by(Document.created_at.desc())
    return list(session.execute(stmt).scalars().all())


def permanently_delete_document(
    session: Session, settings: Settings, document_id: str, *, confirm_filename: str
) -> None:
    document = get_document_or_404(session, document_id)
    if confirm_filename != document.original_filename:
        raise DocumentPermanentDeletionError(
            "Typed confirmation does not match this document's exact original filename."
        )

    try:
        create_backup(settings, category="pre_delete")
    except BackupError as exc:
        raise DocumentPermanentDeletionError(f"Refusing to delete: rollback backup failed: {exc}") from exc

    remove_document_fts(session, document.id)
    remove_recall(session, "document", document.id)

    stored_path = settings.documents_dir / document.stored_relative_path
    stored_path.unlink(missing_ok=True)

    session.delete(document)
    session.commit()
